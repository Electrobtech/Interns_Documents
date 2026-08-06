"""Document text extraction — carried forward, plus the XLSX loader Phase 6 needs.

Each loader's job is to produce *clean text with structure preserved*, because
chunking.py relies on markdown-ish structure (headings, list markers, table
pipes) to keep related content together.
"""
from __future__ import annotations

import csv
import io
import re
from dataclasses import dataclass

import structlog

from orbq_core.errors import ValidationError

log = structlog.get_logger()

SUPPORTED_EXTENSIONS = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".xlsx": "xlsx",
    ".xlsm": "xlsx",
    ".csv": "csv",
    ".tsv": "csv",
    ".txt": "txt",
    ".md": "md",
    ".markdown": "md",
    ".html": "web",
    ".htm": "web",
}

MAX_BYTES = 50 * 1024 * 1024  # 50 MB


@dataclass(slots=True)
class ExtractedDocument:
    text: str
    source_type: str
    metadata: dict


def detect_source_type(filename: str) -> str:
    ext = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    source_type = SUPPORTED_EXTENSIONS.get(ext)
    if not source_type:
        raise ValidationError(
            f"Unsupported file type '{ext or filename}'. Supported: "
            + ", ".join(sorted(SUPPORTED_EXTENSIONS))
        )
    return source_type


# ---------------------------------------------------------------------------
# Per-format extractors
# ---------------------------------------------------------------------------


def _extract_pdf(data: bytes) -> tuple[str, dict]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover
        raise ValidationError("pypdf is required to ingest PDF files") from exc

    reader = PdfReader(io.BytesIO(data))
    pages: list[str] = []
    for i, page in enumerate(reader.pages, 1):
        text = (page.extract_text() or "").strip()
        if text:
            # Page markers double as heading breadcrumbs, so a citation can say
            # which page a claim came from.
            pages.append(f"## Page {i}\n\n{text}")

    if not pages:
        raise ValidationError(
            "No extractable text found. This PDF is likely scanned images — "
            "OCR is required before it can be ingested."
        )
    return "\n\n".join(pages), {"page_count": len(reader.pages)}


def _extract_docx(data: bytes) -> tuple[str, dict]:
    try:
        import docx
    except ImportError as exc:  # pragma: no cover
        raise ValidationError("python-docx is required to ingest DOCX files") from exc

    document = docx.Document(io.BytesIO(data))
    parts: list[str] = []

    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name or "").lower()
        if style.startswith("heading"):
            level = "".join(c for c in style if c.isdigit()) or "1"
            parts.append(f"{'#' * min(int(level), 6)} {text}")
        elif style.startswith("list"):
            parts.append(f"- {text}")
        else:
            parts.append(text)

    # Tables rendered as markdown pipes so chunking keeps rows atomic.
    for table in document.tables:
        rows = [
            "| " + " | ".join(cell.text.strip() for cell in row.cells) + " |"
            for row in table.rows
        ]
        if rows:
            parts.append("\n".join(rows))

    return "\n\n".join(parts), {"paragraphs": len(document.paragraphs), "tables": len(document.tables)}


def _extract_xlsx(data: bytes) -> tuple[str, dict]:
    """Phase 6 addition. Sheet-aware and header-preserving.

    Repeating the header on every row block is deliberate: a retrieved chunk of
    bare cells is meaningless without knowing what the columns are.
    """
    try:
        from openpyxl import load_workbook
    except ImportError as exc:  # pragma: no cover
        raise ValidationError("openpyxl is required to ingest XLSX files") from exc

    workbook = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts: list[str] = []
    total_rows = 0

    for sheet in workbook.worksheets:
        rows = list(sheet.iter_rows(values_only=True))
        if not rows:
            continue
        parts.append(f"## Sheet: {sheet.title}")

        header = [str(c) if c is not None else "" for c in rows[0]]
        body = rows[1:]
        total_rows += len(body)

        parts.append("| " + " | ".join(header) + " |")
        parts.append("|" + "---|" * len(header))

        for row in body:
            cells = ["" if c is None else str(c) for c in row]
            if any(cells):
                parts.append("| " + " | ".join(cells) + " |")

    workbook.close()
    if not parts:
        raise ValidationError("The spreadsheet contains no readable rows")
    return "\n".join(parts), {"sheets": len(workbook.worksheets), "rows": total_rows}


def _extract_csv(data: bytes) -> tuple[str, dict]:
    text = data.decode("utf-8", errors="replace")
    dialect_delim = "\t" if "\t" in text.split("\n", 1)[0] else ","
    reader = csv.reader(io.StringIO(text), delimiter=dialect_delim)

    rows = list(reader)
    if not rows:
        raise ValidationError("The CSV file is empty")

    header, body = rows[0], rows[1:]
    lines = ["| " + " | ".join(header) + " |", "|" + "---|" * len(header)]
    lines.extend("| " + " | ".join(r) + " |" for r in body if any(r))
    return "\n".join(lines), {"rows": len(body), "columns": len(header)}


def _extract_html(data: bytes) -> tuple[str, dict]:
    raw = data.decode("utf-8", errors="replace")

    # Script/style content is not prose and would pollute both the embedding
    # and the keyword index.
    raw = re.sub(r"<(script|style)[^>]*>.*?</\1>", " ", raw, flags=re.DOTALL | re.I)

    for level in range(1, 7):
        raw = re.sub(
            rf"<h{level}[^>]*>(.*?)</h{level}>",
            lambda m, lv=level: f"\n\n{'#' * lv} {m.group(1)}\n\n",
            raw,
            flags=re.DOTALL | re.I,
        )
    raw = re.sub(r"<li[^>]*>(.*?)</li>", r"\n- \1", raw, flags=re.DOTALL | re.I)
    raw = re.sub(r"</(p|div|tr|br)[^>]*>", "\n", raw, flags=re.I)
    raw = re.sub(r"<[^>]+>", " ", raw)

    import html as html_lib

    text = html_lib.unescape(raw)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    return text, {}


def _extract_text(data: bytes) -> tuple[str, dict]:
    return data.decode("utf-8", errors="replace"), {}


_DISPATCH = {
    "pdf": _extract_pdf,
    "docx": _extract_docx,
    "xlsx": _extract_xlsx,
    "csv": _extract_csv,
    "web": _extract_html,
    "txt": _extract_text,
    "md": _extract_text,
    "note": _extract_text,
}


def extract(data: bytes, *, filename: str, source_type: str | None = None) -> ExtractedDocument:
    if len(data) > MAX_BYTES:
        raise ValidationError(
            f"File is {len(data) // 1024 // 1024} MB; the limit is {MAX_BYTES // 1024 // 1024} MB"
        )
    if not data:
        raise ValidationError("File is empty")

    stype = source_type or detect_source_type(filename)
    extractor = _DISPATCH.get(stype)
    if extractor is None:
        raise ValidationError(f"No extractor registered for '{stype}'")

    text, metadata = extractor(data)
    text = text.strip()
    if not text:
        raise ValidationError("No text could be extracted from this file")

    log.info("document_extracted", filename=filename, source_type=stype, chars=len(text))
    return ExtractedDocument(
        text=text,
        source_type=stype,
        metadata=metadata | {"filename": filename, "byte_size": len(data)},
    )
